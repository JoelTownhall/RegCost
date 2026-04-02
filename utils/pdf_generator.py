"""
Document generation for the IALA Policy Toolkit.
- generate_ia_draft_docx(): Word doc IA draft from AI conversation
- generate_rbe_report_pdf(): PDF Regulatory Burden Estimate report
"""
from io import BytesIO
from datetime import date


def generate_ia_draft_docx(proposal: str, draft_text: str) -> bytes:
    """
    Generate a Word document from an AI-compiled IA draft.

    Args:
        proposal:   The user's policy proposal description.
        draft_text: The AI-compiled structured draft (markdown-ish text).

    Returns:
        bytes: The .docx file content.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_heading("Impact Analysis — Working Draft", 0)
    title_para.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Metadata
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("Status: Draft — prepared with IALA Policy Toolkit AI assistance")
    doc.add_paragraph(
        "Note: This document was drafted with AI assistance and requires review before submission to IALA."
    ).runs[0].italic = True

    doc.add_paragraph("")  # spacer

    # Proposal summary
    doc.add_heading("Policy Proposal", 1)
    doc.add_paragraph(proposal or "[No proposal description provided]")

    doc.add_paragraph("")

    # Parse and add the AI-compiled draft sections
    _parse_draft_into_doc(doc, draft_text)

    # Footer note
    doc.add_paragraph("")
    doc.add_paragraph("─" * 60)
    doc.add_paragraph(
        "Prepared using the IALA Policy Toolkit (Impact Analysis Lord Admiralty). "
        "Contact IALA: admiral@iala.ahoy | 1800-AHOY-IALA"
    ).runs[0].font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _parse_draft_into_doc(doc, text: str):
    """
    Parse the AI's markdown-ish draft text into Word document elements.
    Handles **bold headings**, bullet points, and plain paragraphs.
    """
    from docx.shared import Pt, RGBColor

    if not text:
        doc.add_paragraph("[Draft not available]")
        return

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Markdown headings
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("**Question ") and line.endswith("**"):
            # Bold question heading style: **Question N: ...**
            heading_text = line.strip("*").strip()
            h = doc.add_heading(heading_text, 1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif line.startswith("**") and line.endswith("**") and len(line) > 4:
            # Bold line — treat as subheading
            doc.add_heading(line.strip("*").strip(), 2)
        elif line.startswith("- ") or line.startswith("* "):
            # Bullet point
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("  - ") or line.startswith("  * "):
            # Nested bullet
            doc.add_paragraph(line[4:].strip(), style="List Bullet 2")
        elif line.strip() == "" or line.strip() == "---":
            # Blank line or divider — add paragraph spacer
            if i > 0 and lines[i - 1].strip():
                doc.add_paragraph("")
        else:
            # Plain paragraph — handle inline **bold**
            para = doc.add_paragraph()
            _add_inline_formatted(para, line)

        i += 1


def _add_inline_formatted(para, text: str):
    """Add a paragraph run, converting **bold** markers to bold runs."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


# ---------------------------------------------------------------------------
# RBE Report PDF (reportlab)
# ---------------------------------------------------------------------------

def generate_rbe_report_pdf(
    proposal: dict,
    admin_df,
    substantive_df,
    delay_df,
    rbe_summary: dict,
    costing_period: int = 10,
) -> bytes:
    """
    Generate a PDF Regulatory Burden Estimate report.

    Args:
        proposal:        dict with keys: title, agency, contact_name, contact_email, date, description
        admin_df:        DataFrame of administrative cost rows (with calculated totals)
        substantive_df:  DataFrame of substantive compliance cost rows
        delay_df:        DataFrame of delay cost rows
        rbe_summary:     dict: {category: {Business, Individuals, Community organisations, Total}}
        costing_period:  int, default 10 years

    Returns:
        bytes: PDF file content
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle,
            Paragraph, Spacer, HRFlowable,
        )
        from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    except ImportError:
        raise RuntimeError("reportlab not installed. Run: pip install reportlab")

    import pandas as pd

    NAVY = colors.HexColor("#1f4e79")
    LIGHT_BLUE = colors.HexColor("#dce8f5")
    MID_BLUE = colors.HexColor("#2e75b6")
    GREY = colors.HexColor("#f5f5f5")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    styles = getSampleStyleSheet()

    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=NAVY, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=MID_BLUE, spaceAfter=4)
    normal = styles["Normal"]
    small = ParagraphStyle("Small", parent=normal, fontSize=8, textColor=colors.grey)
    bold = ParagraphStyle("Bold", parent=normal, fontName="Helvetica-Bold")
    centered = ParagraphStyle("Centered", parent=normal, alignment=TA_CENTER)

    story = []

    # --- Header ---
    story.append(Paragraph("Regulatory Burden Estimate", h1))
    story.append(Paragraph(proposal.get("title") or "Untitled Proposal", h2))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY))
    story.append(Spacer(1, 0.4 * cm))

    # --- Proposal metadata table ---
    meta_data = [
        ["Agency / Department", proposal.get("agency", "")],
        ["Contact officer", proposal.get("contact_name", "")],
        ["Contact email", proposal.get("contact_email", "")],
        ["Date", proposal.get("date", str(date.today()))],
        ["Costing period", f"{costing_period} years"],
        ["Status", "Draft — prepared with IALA Policy Toolkit"],
    ]
    meta_table = Table(meta_data, colWidths=[4.5 * cm, 12 * cm])
    meta_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, GREY]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(meta_table)

    if proposal.get("description"):
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("Description", bold))
        story.append(Paragraph(proposal["description"], normal))

    story.append(Spacer(1, 0.5 * cm))

    # --- RBE Summary Table ---
    story.append(Paragraph("Regulatory Burden Estimate (RBE) Table", h1))
    story.append(Paragraph(
        f"Average annual regulatory costs over {costing_period}-year costing period ($ millions, real terms)",
        small,
    ))
    story.append(Spacer(1, 0.3 * cm))

    cats = ["Administrative costs", "Substantive compliance costs", "Delay costs"]
    stakeholders = ["Business", "Individuals", "Community organisations", "Total"]

    header_row = ["Cost category"] + stakeholders
    rbe_rows = [header_row]
    col_totals = {s: 0.0 for s in stakeholders}

    for cat in cats:
        row = [cat]
        cat_data = rbe_summary.get(cat, {})
        for s in stakeholders:
            val = cat_data.get(s, 0.0) or 0.0
            if s != "Total":
                col_totals[s] += val
            row.append(f"${val:.3f}m" if val else "–")
        rbe_rows.append(row)

    # Grand total row
    grand_total = sum(col_totals[s] for s in ["Business", "Individuals", "Community organisations"])
    col_totals["Total"] = grand_total
    total_row = ["Total"]
    for s in stakeholders:
        v = col_totals[s]
        total_row.append(f"${v:.3f}m" if v else "–")
    rbe_rows.append(total_row)

    col_widths = [6 * cm, 3 * cm, 3 * cm, 3.5 * cm, 2.5 * cm]
    rbe_table = Table(rbe_rows, colWidths=col_widths, repeatRows=1)
    rbe_table.setStyle(TableStyle([
        # Header
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        # Body
        ("FONTSIZE", (0, 1), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, LIGHT_BLUE]),
        # Total row
        ("BACKGROUND", (0, -1), (-1, -1), MID_BLUE),
        ("TEXTCOLOR", (0, -1), (-1, -1), colors.white),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        # Grid
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(rbe_table)
    story.append(Spacer(1, 0.5 * cm))

    # --- Cost detail appendices ---
    def _df_to_table(df, title):
        if df is None or (hasattr(df, "empty") and df.empty):
            return
        story.append(Paragraph(title, h2))
        rows = [list(df.columns)]
        for _, r in df.iterrows():
            rows.append([str(v) if v is not None and str(v) != "nan" else "–" for v in r])
        if len(rows) == 1:
            story.append(Paragraph("No entries recorded.", normal))
            return
        # Distribute columns evenly
        page_width = 17 * cm
        n_cols = len(rows[0])
        col_w = [page_width / n_cols] * n_cols
        t = Table(rows, colWidths=col_w, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, GREY]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ("WORDWRAP", (0, 0), (-1, -1), True),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("Appendix: Cost Input Detail", h1))
    _df_to_table(admin_df, "Administrative Costs")
    _df_to_table(substantive_df, "Substantive Compliance Costs")
    _df_to_table(delay_df, "Delay Costs")

    # --- Methodology notes ---
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Methodology Notes", h1))
    methodology_notes = [
        f"Costing period: {costing_period} years (default IALA period)",
        "Default hourly tariff (work-related): $85.17/hour",
        "Default hourly tariff (leisure/non-employment): $37/hour",
        "Costs presented as average annual in real terms (constant prices)",
        "Administrative costs: costs of demonstrating compliance (form-filling, reporting, inspections, etc.)",
        "Substantive compliance costs: real changes to operations (capital, labour, process changes)",
        "Delay costs: costs from government-imposed waiting periods",
        "Excluded from RBMF: taxes, pure opportunity costs, government enforcement costs, international obligation costs",
    ]
    for note in methodology_notes:
        story.append(Paragraph(f"• {note}", small))

    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        f"Prepared using the IALA Policy Toolkit (Impact Analysis Lord Admiralty) | "
        f"{date.today().strftime('%d %B %Y')} | admiral@iala.ahoy | 1800-AHOY-IALA",
        small,
    ))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()
